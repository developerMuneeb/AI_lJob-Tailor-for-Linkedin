import streamlit as st
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
import spacy
from docx import Document
import re
import os
import pdfplumber

#Load Spcy Model
import spacy.cli
spacy.cli.download("en_core_web_sm")  # <--- this line downloads the model
import spacy
nlp = spacy.load("en_core_web_sm")

# Download NLTK resources
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)

# Load spaCy model
nlp = spacy.load('en_core_web_sm')

# Streamlit app layout
st.title("🎯 LinkedIn Job Tailor: CV & Cover Letter Generator")
st.write("Upload your CV (PDF) and paste a LinkedIn job description to get a tailored CV and cover letter.")

# File uploader for CV
cv_file = st.file_uploader("📄 Upload your CV (PDF)", type=["pdf"])

# Text area for job description
job_description = st.text_area("📋 Paste the LinkedIn Job Description", height=300)

# Button to process
if st.button("🚀 Tailor CV and Generate Cover Letter"):
    if cv_file and job_description:
        with st.spinner("Processing..."):

            # Save uploaded CV temporarily
            with open("temp_cv.pdf", "wb") as f:
                f.write(cv_file.getbuffer())

            # Step 1: Extract text from PDF
            try:
                with pdfplumber.open("temp_cv.pdf") as pdf:
                    text = "".join(page.extract_text() or "" for page in pdf.pages)
                st.success("✅ CV parsed successfully!")
            except Exception as e:
                st.error(f"❌ Error parsing CV: {e}")
                os.remove("temp_cv.pdf")
                st.stop()

            # Step 2: Extract CV Skills using noun chunks
            doc = nlp(text)
            cv_skills = list(set(chunk.text.lower() for chunk in doc.noun_chunks if len(chunk.text) > 2))
            cv_data = {'name': 'Your Name', 'skills': cv_skills}
            st.write("**📌 Extracted Skills from CV:**")
            st.write(cv_skills[:20])

            # Step 3: Extract keywords from job description
            def extract_keywords(text):
                tokens = word_tokenize(text.lower())
                stop_words = set(stopwords.words('english'))
                lemmatizer = WordNetLemmatizer()
                keywords = [lemmatizer.lemmatize(token) for token in tokens
                            if token.isalnum() and token not in stop_words]
                doc = nlp(text)
                entities = [ent.text.lower() for ent in doc.ents]
                noun_chunks = [chunk.text.lower() for chunk in doc.noun_chunks]
                all_keywords = set(keywords + entities + noun_chunks)
                filtered_keywords = [kw for kw in all_keywords if len(kw) > 2 and kw not in ['job', 'role', 'team']]
                return filtered_keywords

            job_keywords = extract_keywords(job_description)
            st.write("**💼 Extracted Keywords from Job Description:**")
            st.write(list(job_keywords)[:20])

            # Step 4: Identify new skills to add
            new_skills = [kw for kw in job_keywords if kw not in [skill.lower() for skill in cv_skills]]
            st.write("**➕ Skills to Consider Adding to CV:**")
            st.write(new_skills[:10] if new_skills else "Your CV already covers most job keywords!")

            # Step 5: Modify CV with additional skills
            skills_section = "\nSkills\n" + ", ".join(cv_skills + new_skills[:10])
            modified_cv = text + "\n\n" + skills_section
            docx_doc = Document()
            docx_doc.add_heading('Modified CV', 0)
            docx_doc.add_paragraph(modified_cv)
            docx_doc.save("modified_cv.docx")
            st.write("**📝 Modified CV Preview:**")
            st.write(modified_cv[:500] + "...")
            with open("modified_cv.docx", "rb") as f:
                st.download_button("⬇️ Download Modified CV", f, "modified_cv.docx")

            # Step 6: Generate Cover Letter
            def generate_cover_letter(cv_data, job_keywords, job_description):
                name = cv_data.get('name', 'Your Name')
                company = re.search(r'company[:\s]+([A-Za-z0-9&.\- ]+)', job_description, re.IGNORECASE)
                company_name = company.group(1).strip() if company else 'Hiring Manager'
                template = f"""
Dear {company_name},

I am excited to apply for the position advertised on LinkedIn. With my background in {', '.join(cv_skills[:3])}, I am confident in my ability to contribute to your team.

The job description highlights the need for skills such as {', '.join(job_keywords[:5])}. In my previous roles, I have demonstrated proficiency in these areas, including {', '.join(new_skills[:2]) if new_skills else 'relevant skills'}. My experience aligns with your requirements, and I am eager to bring my expertise to {company_name}.

Thank you for considering my application. I look forward to the opportunity to discuss how I can contribute to your team.

Sincerely,  
{name}
"""
                return template.strip()

            cover_letter = generate_cover_letter(cv_data, job_keywords, job_description)
            st.write("**📄 Generated Cover Letter:**")
            st.code(cover_letter)
            cl_doc = Document()
            cl_doc.add_heading('Cover Letter', 0)
            cl_doc.add_paragraph(cover_letter)
            cl_doc.save("cover_letter.docx")
            with open("cover_letter.docx", "rb") as f:
                st.download_button("⬇️ Download Cover Letter", f, "cover_letter.docx")

            # Step 7: Cleanup
            for file in ["temp_cv.pdf", "modified_cv.docx", "cover_letter.docx"]:
                if os.path.exists(file):
                    os.remove(file)

    else:
        st.error("Please upload a CV and paste the job description.")

# Footer
st.markdown("---")
st.caption("Built with ❤️ using Streamlit. Make sure your CV is in PDF format and description is from LinkedIn.")
